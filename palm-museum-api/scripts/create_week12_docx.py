from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = Document()

# ─── 样式 ───
style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ─── 标题 ───
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('第12周小组会议纪要')
run.bold = True
run.font.size = Pt(18)

# 基本信息
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
info.add_run('会议日期：').bold = True
info.add_run('2026年5月22日')
info.add_run('\n会议主题：').bold = True
info.add_run('集成测试、Bug修复、演示准备')
info.add_run('\n参会人员：').bold = True
info.add_run('全体小组成员')

doc.add_paragraph('—' * 40)

# ═══════════════════════════════════════
# 一、工作进展回顾
# ═══════════════════════════════════════
doc.add_heading('一、工作进展回顾', level=1)

doc.add_paragraph('第11周所有11项核心功能开发已全部完成，本周进入测试与完善阶段。')

# ═══════════════════════════════════════
# 二、本周完成工作
# ═══════════════════════════════════════
doc.add_heading('二、本周完成工作', level=1)

items = [
    ('图片素材搜集', '搜集25件文物的真实图片并保存至后端images/目录，替换了此前所有example.com占位URL。将图片URL改为相对路径（/images/por_001.jpg），后端返回时自动拼接当前IP，换网络只需改HttpClient.ets一处。'),
    ('以图搜图功能上线', '运行seed_features.py成功提取全部25件文物128维视觉特征入库。经Swagger验证，上传图片可返回按相似度排序的匹配结果，以图搜图功能正式可用。'),
    ('语音导览（TTS）接入', '后端创建TTS API，使用edge-tts（微软中文语音XiaoxiaoNeural）实时合成讲解音频，首次访问后自动缓存。前端AudioPlayer组件支持真机播放和预览器降级两种模式。'),
    ('UI美化', '统一按钮风格（阴影、圆角、颜色），建立三级阴影体系。全部18处Scroll组件添加弹性回弹效果。登录页Logo放大至120px并调整布局，删除"或"分隔线。'),
]

for title, desc in items:
    p = doc.add_paragraph()
    p.add_run(f'{title}：').bold = True
    p.add_run(desc)

# ═══════════════════════════════════════
# 三、Bug修复
# ═══════════════════════════════════════
doc.add_heading('三、Bug修复', level=1)

bugs = [
    '首页空白（后端不可用时Mock数据不加载）——修复：ArtworkService.fetchArtworks()失败时自动回退到MockArtworks数据',
    '模拟器因缺少系统能力无法运行——修复：添加sysCapCheck = false跳过系统能力校验（后因SDK版本不支持改为使用预览器）',
    'IP变更导致前后端连不上——修复：数据库图片URL改为相对路径，客户端用HttpClient.getBaseUrl()拼接',
    '今日推荐卡片按钮文字重叠——修复：调整布局间距',
    '热门搜索页默认展开——修复：改为点击搜索框时弹出',
]
for bug in bugs:
    doc.add_paragraph(bug, style='List Bullet')

# ═══════════════════════════════════════
# 四、决议事项
# ═══════════════════════════════════════
doc.add_heading('四、决议事项', level=1)

resolutions = [
    '项目核心功能已全部完成，进入测试和完善阶段',
    '6个已发现的Bug由对应负责人按期修复',
    '答辩PPT突出以图搜图算法和HarmonyOS ArkTS技术亮点',
    '语音导览暂用TTS方案代替人工录制',
]
for i, res in enumerate(resolutions, 1):
    doc.add_paragraph(f'{i}. {res}')

# ═══════════════════════════════════════
# 五、待办事项
# ═══════════════════════════════════════
doc.add_heading('五、待办事项', level=1)

todos = [
    'Bug修复按期完成',
    '用户手册初稿',
    '答辩PPT准备',
    '功能演示视频录制',
]
for item in todos:
    doc.add_paragraph(f'☐ {item}')

# ─── 尾部 ───
doc.add_paragraph('—' * 40)
p = doc.add_paragraph()
p.add_run('下次会议时间：').bold = True
p.add_run('2026年5月27日')
p.add_run('\n下次会议主题：').bold = True
p.add_run('答辩前最终检查与预演')

output_path = '/sessions/quirky-blissful-tesla/mnt/museum/文本/第12周_5.22_集成测试与Bug修复.docx'
doc.save(output_path)
print(f'Document saved to: {output_path}')
