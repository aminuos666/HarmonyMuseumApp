from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Arial'
font.size = Pt(11)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('第12周小组会议纪要')
run.bold = True
run.font.size = Pt(18)

info = doc.add_paragraph()
info.add_run('会议日期：').bold = True
info.add_run('2026年5月22日')
info.add_run('\n会议主题：').bold = True
info.add_run('集成测试、Bug修复、演示准备')
info.add_run('\n参会人员：').bold = True
info.add_run('全体小组成员')

doc.add_paragraph('—' * 40)

doc.add_heading('一、工作进展回顾', level=1)
doc.add_paragraph('第11周计划中的11项核心功能开发已全部完成，项目进入测试与完善阶段。')

doc.add_heading('二、本周完成工作', level=1)

p = doc.add_paragraph()
p.add_run('以图搜图功能完善：').bold = True
p.add_run('完成25件文物真实图片的搜集与入库，运行seed_features.py提取全部128维视觉特征。数据库图片URL改为相对路径（/images/por_001.jpg），后端返回时自动拼接当前IP，换网络只需修改HttpClient.ets一处。')

p = doc.add_paragraph()
p.add_run('语音导览（TTS）接入：').bold = True
p.add_run('后端创建TTS API，使用edge-tts实现中文语音合成。前端AudioPlayer支持真机播放和预览器降级两种模式。')

p = doc.add_paragraph()
p.add_run('UI美化：').bold = True
p.add_run('统一按钮样式（阴影、圆角），建立三级阴影体系。全部页面Scroll组件添加弹性回弹效果。登录页Logo放大，删除"或"分隔线。')

p = doc.add_paragraph()
p.add_run('Bug修复：').bold = True
p.add_run('后端不可用时首页空白（Mock数据未加载）——修复：失败时自动回退Mock数据。图片IP硬编码导致换网络显示不了——修复：改为相对路径+动态拼接。今日推荐按钮文字重叠——修复：调整间距。搜索页热门搜索常驻——修复：改为点击搜索框弹出。')

doc.add_heading('三、决议事项', level=1)
doc.add_paragraph('1. 全部核心功能已完成，进入测试完善阶段')
doc.add_paragraph('2. 6个已发现Bug由负责人按期修复')
doc.add_paragraph('3. 答辩PPT突出以图搜图算法和ArkTS技术亮点')
doc.add_paragraph('4. 语音导览暂用TTS替代人工录制')

doc.add_heading('四、待办事项', level=1)
doc.add_paragraph('☐ Bug修复')
doc.add_paragraph('☐ 用户手册初稿')
doc.add_paragraph('☐ 答辩PPT准备')
doc.add_paragraph('☐ 功能演示视频录制')

doc.add_paragraph('—' * 40)
p = doc.add_paragraph()
p.add_run('下次会议时间：').bold = True
p.add_run('2026年5月27日')
p.add_run('\n下次会议主题：').bold = True
p.add_run('答辩前最终检查与预演')

doc.save('/sessions/quirky-blissful-tesla/mnt/outputs/第12周_5.22_集成测试与Bug修复.docx')
print('done')
